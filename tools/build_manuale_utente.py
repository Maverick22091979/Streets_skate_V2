from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/manuale_utente_street_skate.docx")
FALLBACK_OUT = Path("artifacts/manuale_utente_street_skate_updated.docx")
LOGIN_IMG = Path("artifacts/manual_images/login.png")
DASHBOARD_IMG = Path(r"C:\Users\giorg\Pictures\Screenshots\Screenshot 2026-07-26 143332.png")


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_table_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("475467")


def style_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Heading 1", 16, "163A63"),
        ("Heading 2", 13, "163A63"),
        ("Heading 3", 11.5, "345D8C"),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    if "ManualCallout" not in styles:
        s = styles.add_style("ManualCallout", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.font.name = "Calibri"
        s.font.size = Pt(10.5)
        s.font.color.rgb = RGBColor.from_string("0F172A")
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(6)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Manuale Utente Street Skate")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "Guida operativa basata sul codice applicativo corrente. Copre accesso utente, accesso admin, import, filtri, dettagli percorso, controlli duplicati e impostazioni."
    )
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_table(rows=3, cols=2)
    set_table_layout(meta, [1.8, 4.7])
    rows = [
        ("Versione documento", "1.3"),
        ("Data", date(2026, 8, 21).strftime("%d/%m/%Y")),
        ("Ambito", "Web app FastAPI Street Skate con provider Strava, adidas Running, MapMyRun, Suunto, Runkeeper, Garmin, Inline Route Tracking e accesso Admin"),
    ]
    for i, (a, b) in enumerate(rows):
        set_cell(meta.rows[i].cells[0], a, bold=True)
        shade_cell(meta.rows[i].cells[0], "E8EEF5")
        set_cell(meta.rows[i].cells[1], b)


def add_overview(doc: Document):
    doc.add_heading("1. Panoramica applicativa", level=1)
    doc.add_paragraph(
        "Street Skate è un'applicazione web per importare percorsi sportivi, arricchirli con dati derivati e consultarli tramite dashboard. Il sistema salva i percorsi su PostgreSQL e supporta due profili operativi: utente loggato tramite provider sportivo e admin locale."
    )

    tbl = doc.add_table(rows=1, cols=4)
    set_table_layout(tbl, [1.25, 1.6, 2.15, 1.5])
    for c, t in zip(tbl.rows[0].cells, ["Area", "Modalità", "Descrizione", "Disponibilità"]):
        set_cell(c, t, bold=True)
        shade_cell(c, "DCE6F2")
    rows = [
        ("Login", "Provider OAuth/export", "Accesso con account sportivo o export adidas locale", "Utente"),
        ("Dashboard", "Percorsi", "Vista elenco, filtri, import, delete, riepiloghi", "Utente/Admin"),
        ("Dettaglio", "Percorso", "Mappa Leaflet e metriche del singolo percorso", "Utente/Admin"),
        ("Admin", "Impostazioni", "Controlli duplicati, pendenze, audit log, tipologia percorsi e note sui punteggi fissi", "Admin"),
    ]
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)

    doc.add_paragraph(
        "La home espone i pulsanti di accesso per adidas Running, Strava, MapMyRun, Suunto, Runkeeper, Garmin Connect e Inline Route Tracking. Provider non configurati restano non utilizzabili fino alla configurazione delle credenziali o dell'eventuale export previsto."
    )
    doc.add_paragraph(
        "Nel deploy pubblico corrente l'applicazione gira su VPS Aruba con Cloudflare davanti, nginx installato sull'host come reverse proxy HTTPS e container Docker dedicati a web e database."
    )


def add_login_modes(doc: Document):
    doc.add_heading("2. Modalità di accesso", level=1)
    doc.add_heading("2.1 Accesso utente tramite provider", level=2)
    add_bullets(
        doc,
        [
            "Strava: il pulsante apre la pagina ufficiale di login del provider; l'utente inserisce lì le credenziali e il backend completa l'autenticazione OAuth tramite callback e token API.",
            "MapMyRun: se configurato, il pulsante apre la pagina ufficiale di login del provider; l'app non legge username e password ma riceve il token API dopo il consenso OAuth.",
            "Suunto: il provider e predisposto nel progetto, ma al 19 agosto 2026 la richiesta partnership risulta inviata e in review da parte di Suunto, con finestra indicativa di risposta entro circa il 2 settembre 2026.",
            "Runkeeper: il provider e predisposto come integrazione legacy e va verificato sul campo prima di un uso reale.",
            "Garmin Connect: il supporto nel progetto è predisposto, ma al 3 agosto 2026 i servizi API Garmin risultano in migrazione secondo il team developer Garmin e non è stata comunicata una data di ripristino. L'attivazione reale resta quindi bloccata anche in presenza della configurazione locale.",
            "Inline Route Tracking: il provider è già predisposto nella schermata di login rapido ma resta in attesa delle API ufficiali di login/auth e dello schema callback dell'app Android.",
            "adidas Running: il pulsante provider è riservato al login ufficiale OAuth; se l'OAuth non è configurato ma l'export locale è disponibile, l'accesso tramite export compare come azione separata.",
            "Dopo un login utente riuscito, all'apertura della dashboard compare un popup di consenso download percorsi nel database della web app.",
        ],
    )
    doc.add_heading("2.2 Accesso Admin locale", level=2)
    add_bullets(
        doc,
        [
            "L'admin accede con username e password locali definiti nel file .env.",
            "L'admin non usa provider esterni per autenticarsi.",
            "Il login admin applica una protezione anti brute force in memoria: dopo 5 tentativi falliti nella finestra corrente, l'accesso viene bloccato temporaneamente per 15 minuti per combinazione IP/username.",
            "Una volta loggato visualizza tutti i percorsi presenti in archivio, non solo quelli di un singolo utente.",
        ],
    )

    doc.add_heading("2.3 Esito della sessione", level=2)
    add_bullets(
        doc,
        [
            "Per gli utenti provider la sessione salva user_id e connection_id.",
            "Per l'admin la sessione salva il contesto local_admin.",
            "Con la configurazione corrente i login provider aprono l'archivio globale in sola lettura: tutti i percorsi e i relativi punteggi sono visibili, ma import e cancellazioni restano riservati all'admin.",
            "Il logout cancella la sessione locale e, per i provider che supportano una revoca ufficiale via API, tenta anche la chiusura della sessione/token lato provider prima di registrare l'evento nell'audit log.",
            "Le sessioni autenticate registrano timestamp di avvio e ultimo utilizzo; con la configurazione corrente, se l'utente resta inattivo per oltre 3 minuti la sessione viene invalidata e deve rifare login.",
            "Con la configurazione corrente il cookie di sessione è non persistente: alla chiusura del browser la sessione lato client decade e il login va ripetuto alla riapertura.",
            "Dopo il logout, la pagina iniziale riporta i pulsanti provider a un nuovo ciclo di login, così il click successivo riparte dal flusso di autenticazione del provider relativo.",
            "Le credenziali del provider non transitano nell'app Street Skate: vengono inserite nella pagina ufficiale del provider e l'app riceve solo code/token di autenticazione.",
            "Nella schermata di login ogni provider mostra lo stato corrente di implementazione per Login, Auth e Settings Secrets.",
            "La schermata di login non espone più password di configurazione nei placeholder e le variabili sensibili devono essere valorizzate solo tramite file ambiente o secret manager.",
            "Lo sviluppo locale usa il file .env con callback localhost, mentre il deploy VPS usa un file separato .env.production con BASE_URL pubblico e redirect URI coerenti con il dominio pubblicato.",
            "I cookie di sessione sono configurabili via ambiente; in produzione il flag Secure deve restare attivo dietro HTTPS, mentre in sviluppo locale può restare disattivato per localhost.",
            "Le operazioni mutanti via form e via fetch usano un token CSRF di sessione, verificato dal backend prima di accettare login admin, logout, import, cancellazioni e salvataggi amministrativi.",
            "Il backend invia header HTTP di hardening come Content-Security-Policy, X-Content-Type-Options, X-Frame-Options, Referrer-Policy, Permissions-Policy e header cross-origin restrittivi; HSTS va attivato solo nel deploy HTTPS pubblico.",
            "Lo script della pagina dettaglio percorso è stato spostato in un file statico dedicato, così la CSP può evitare unsafe-inline per gli script applicativi.",
        ],
    )


def add_screenshots(doc: Document):
    doc.add_heading("3. Schermate principali", level=1)
    doc.add_paragraph(
        "Le immagini seguenti mostrano la schermata di accesso e un esempio di dashboard percorsi."
    )

    if LOGIN_IMG.exists():
        doc.add_heading("3.1 Login", level=2)
        doc.add_picture(str(LOGIN_IMG), width=Inches(4.4))
        add_caption(doc, "Schermata login con accesso provider e accesso Admin locale.")

    if DASHBOARD_IMG.exists():
        doc.add_heading("3.2 Dashboard percorsi", level=2)
        doc.add_picture(str(DASHBOARD_IMG), width=Inches(6.2))
        add_caption(doc, "Dashboard percorsi con pulsanti import, riepiloghi e filtri laterali.")


def add_user_dashboard(doc: Document):
    doc.add_heading("4. Funzionalità disponibili per utente loggato", level=1)
    doc.add_heading("4.1 Dashboard utente", level=2)
    add_bullets(
        doc,
        [
            "Visualizzazione del nome utente/provider corrente nell'header.",
            "Al login provider, la dashboard carica subito dal database l'archivio globale dei percorsi, non solo quelli dell'utente autenticato.",
            "Il pulsante Importa risulta disabilitato per i login provider, perché il catalogo globale e in sola lettura.",
            "Pulsante Ricarica cache locale per rileggere l'archivio globale gia salvato nel database.",
            "Barra di avanzamento che indica i percorsi trovati nell'archivio globale.",
            "Cinque riquadri riepilogativi: numero percorsi, distanza totale, difficoltà media, percorsi difficili, ultima importazione.",
        ],
    )
    doc.add_heading("4.2 Filtri percorsi", level=2)
    add_bullets(
        doc,
        [
            "Filtro per difficoltà dichiarata iniziale derivata dal calendario eventi.",
            "Filtro per Utente importatore, disponibile anche nei login provider in sola lettura.",
            "Distanza km massima.",
            "Dislivello m massimo.",
            "Velocità massima km/h massima.",
            "Velocità media km/h massima.",
            "Pendenza massima assoluta % massima.",
            "Pendenza media ponderata assoluta % massima.",
            "Filtro Sampietrini sì/no derivato dall'enrichment OSM.",
            "Pendenza media % massima.",
            "Score difficoltà massimo.",
            "Reset filtri per tornare all'elenco completo visibile all'utente.",
        ],
    )
    doc.add_heading("4.3 Elenco percorsi", level=2)
    add_bullets(
        doc,
        [
            "Ogni card mostra score e livello di difficoltà.",
            "Accanto al nome del percorso compare un indicatore circolare con il colore associato alla difficoltà dichiarata iniziale.",
            "Sono presenti nome percorso, tipo attività, data locale, distanza, dislivello, velocità media, pendenza massima assoluta, verso della pendenza massima, superficie, presenza sampietrini, pressione, temperatura e fonte dei dati.",
            "Ogni card include il nome dell'utente importatore registrato sul record.",
            "Nei login provider l'elenco e in sola lettura e non mostra il pulsante Cancella.",
            "Il link sul nome apre il dettaglio del percorso.",
        ],
    )
    doc.add_heading("4.4 Dettaglio percorso", level=2)
    add_bullets(
        doc,
        [
            "Mappa Leaflet con tracciato, marker di partenza e arrivo.",
            "Statistiche principali: distanza, dislivello, pendenza media, velocità media, difficoltà, data di importazione.",
            "Fallback su bounding box o vista iniziale se i punti del tracciato non sono disponibili.",
        ],
    )


def add_import_logic(doc: Document):
    doc.add_heading("5. Logica di import e regole di salvataggio", level=1)
    doc.add_paragraph(
        "Le regole di import determinano quali percorsi vengono scartati e quali vengono scritti nel database. La logica è condivisa tra import provider e import manuale, con differenze specifiche per il ruolo."
    )
    doc.add_heading("5.1 Sorgenti supportate", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    set_table_layout(tbl, [1.45, 1.55, 2.0, 1.5])
    for c, t in zip(tbl.rows[0].cells, ["Provider", "Metodo", "Dati importati", "Utente"]):
        set_cell(c, t, bold=True)
        shade_cell(c, "DCE6F2")
    rows = [
        ("Strava", "OAuth", "Attività con polyline e metriche base", "Utente"),
        ("MapMyRun", "OAuth", "Workout con time series e punti", "Utente"),
        ("adidas Running", "OAuth/export", "Attività API o file GPX/JSON da export", "Utente"),
        ("Inline Route Tracking", "OAuth futuro", "Provider predisposto, in attesa API ufficiali", "Utente"),
        ("Manuale GPX", "Upload file", "GPX e sidecar JSON", "Admin"),
    ]
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)

    doc.add_heading("5.1.1 Consenso download percorsi", level=3)
    add_bullets(
        doc,
        [
            "Dopo autenticazione utente riuscita, Street Skate mostra un popup/disclaimer in overlay sulla dashboard.",
            "Il disclaimer chiede l'autorizzazione a scaricare e salvare nel DB i percorsi dell'utente.",
            "Se l'utente accetta, l'import provider parte automaticamente in background e il footer mostra un conteggio realtime nel formato [percorsi importati / percorsi totali]. A completamento import, il messaggio finale resta visibile per 8 secondi e poi il footer si nasconde.",
            "Se l'utente rifiuta, l'accesso alla dashboard è consentito ma l'import viene bloccato lato interfaccia e lato API.",
        ],
    )

    doc.add_heading("5.2 Criteri di scarto comuni", level=2)
    add_bullets(
        doc,
        [
            "Percorsi sotto la distanza minima configurata (MIN_IMPORT_DISTANCE_M).",
            "Percorsi senza external_id valido nel caso di provider API.",
            "Percorsi considerati duplicati secondo le soglie definite dall'admin: differenza percentuale massima di distanza, tolleranza sugli estremi del tracciato e controllo del percorso invertito.",
            "Con la logica corrente il nome percorso viene usato come primo indicatore, ma la decisione finale di duplicato rispetta i parametri di similarità configurati.",
        ],
    )
    doc.add_heading("5.3 Arricchimento automatico", level=2)
    add_bullets(
        doc,
        [
            "Calcolo profilo pendenza e pendenza media ponderata.",
            "Recupero superficie e smoothness da OpenStreetMap/Overpass, se disponibile.",
            "Derivazione del flag has_cobblestone dai campi surface OSM, con rilevazione dei valori cobblestone o sett.",
            "Recupero meteo storico o forecast da Open-Meteo, se disponibile.",
            "Calcolo indice difficoltà Street Skate basato sui pesi base dell'algoritmo.",
            "Per i file GPX con data riconoscibile nel nome, generazione di una difficoltà dichiarata iniziale da calendario eventi (Easy, EasyLong, Advanced, Pro).",
            "Registrazione di import_user_label per sapere quale utente ha eseguito il download/import.",
        ],
    )

    doc.add_heading("5.4 Differenze struttura dati provider", level=2)
    add_bullets(
        doc,
        [
            "La baseline logica Street Skate usa gli stessi campi principali di adidas Running: external_id, name, sport_type, start_date_local, distance_m, elevation_gain_m, average_speed_ms, moving_time_sec, polyline/punti e raw_payload.",
            "Strava: fornisce quasi tutti i campi base ma nel flusso attuale importa soprattutto summary_polyline e metriche aggregate, non un sidecar locale equivalente all'export adidas.",
            "MapMyRun: usa una struttura diversa con aggregates e time_series.position; Street Skate converte i campi compatibili e usa i punti GPS al posto della polyline quando necessario.",
            "Garmin Connect: il CSV esportato manualmente non contiene necessariamente tutti i punti. Con accesso ufficiale all'Activity API Garmin è possibile recuperare file completi FIT/GPX/TCX e dettagli attività più ricchi.",
            "Inline Route Tracking: il provider è stato predisposto con mapping tecnico iniziale, ma i campi reali andranno confermati quando Daniele pubblicherà le API ufficiali di login, profilo e attività.",
            "Quando un provider non espone un campo presente in adidas Running, Street Skate importa solo il sottoinsieme compatibile e segnala il contesto all'utente nella pagina di consenso.",
        ],
    )


def add_admin(doc: Document):
    doc.add_heading("6. Funzionalità disponibili per Admin", level=1)
    doc.add_heading("6.1 Dashboard admin", level=2)
    add_bullets(
        doc,
        [
            "Visualizza l'intero archivio percorsi presente nel database.",
            "Il pulsante Importa da provider è disabilitato.",
            "Può cancellare qualsiasi percorso dalla dashboard.",
            "Ha a disposizione un filtro a discesa per Utente importatore che mostra solo utenti con almeno un percorso importato.",
            "Può aprire la pagina Tipologia percorsi per cambiare il tipo attività salvato su ogni record.",
            "Può aprire la pagina Debug richiesta per verificare client host e header HTTP ricevuti da FastAPI durante accessi via Tailscale Funnel o browser locali.",
        ],
    )
    doc.add_heading("6.2 Import manuale GPX + JSON", level=2)
    add_bullets(
        doc,
        [
            "Upload multiplo di file .gpx e .json sidecar.",
            "Una coppia GPX + JSON con lo stesso nome base viene trattata come un solo bundle e genera un solo percorso arricchito.",
            "Barra di avanzamento dedicata per il batch manuale.",
            "Scarto per file troppo corti, tipo attività non ammesso, nome file già presente o contenuto ritenuto duplicato.",
            "La deduplica dell'import manuale confronta il nuovo percorso contro tutta la cache già presente, non solo contro i percorsi manuali.",
            "I percorsi manuali vengono assegnati al contesto utente admin locale.",
        ],
    )
    doc.add_heading("6.3 Impostazioni amministrative", level=2)
    add_bullets(
        doc,
        [
            "Punteggi Street Skate: i pesi dell'algoritmo base non sono modificabili dalla dashboard admin e il punteggio resta fisso una volta salvato sul percorso.",
            "Pulsanti autenticazione provider: pannello admin con checkbox per decidere se ciascun provider OAuth deve essere visibile nella home e se il relativo pulsante deve risultare cliccabile.",
            "Controllo duplicati GPX: differenza distanza totale max, tolleranza estremi tracciato, controllo percorso invertito.",
            "Calcolo pendenze: finestra smoothing quota, lunghezza minima tratto, cap pendenza massima.",
            "Tipologia percorsi: tabella admin con ricerca testuale e salvataggio della tipologia per singolo percorso.",
            "Audit log: elenco delle azioni registrate dal sistema, con IP origine, metodo/path richiesta, user-agent e geolocalizzazione approssimata derivata dall'IP nello stesso archivio audit_logs.",
            "Debug richiesta: pagina tecnica per confrontare request.client e header di forwarding (es. X-Forwarded-For, Forwarded, X-Real-IP) durante i test di pubblicazione.",
        ],
    )


def add_admin_settings_detail(doc: Document):
    doc.add_heading("7. Dettaglio impostazioni amministrative", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_layout(tbl, [2.2, 2.2, 2.1])
    for c, t in zip(tbl.rows[0].cells, ["Gruppo", "Campi", "Effetto"]):
        set_cell(c, t, bold=True)
        shade_cell(c, "DCE6F2")
    rows = [
        ("Punteggi Street Skate", "Algoritmo base", "Il calcolo usa pesi fissi e non modificabili da interfaccia"),
        ("Pulsanti provider", "visible, enabled per provider", "Controlla visibilità e click dei pulsanti OAuth nella pagina login"),
        ("Controllo duplicati", "distance_diff_pct, endpoint_tolerance_m, allow_reverse_match", "Definiscono quando due tracciati sono considerati lo stesso percorso"),
        ("Calcolo pendenze", "smoothing_window, min_run_distance_m, max_cap_pct", "Influiscono sul profilo altimetrico e sui valori di pendenza"),
        ("Tipologia percorsi", "sport_type", "Permette all'admin di correggere il tipo attività del percorso salvato"),
        ("Audit log", "limit, refresh, IP, geo, request metadata", "Permette consultazione operativa degli eventi con tracciamento origine client"),
        ("Viewer", "show_direction_arrows via API admin", "Imposta il comportamento del visualizzatore percorso"),
    ]
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)

    doc.add_paragraph(
        "Le impostazioni vengono salvate tramite endpoint amministrativi e restano in memoria dell'applicazione in esecuzione. Per i parametri di deduplica, il valore di default attuale è 5% di differenza massima distanza e 300 m di tolleranza estremi. I punteggi difficoltà, invece, usano pesi base e non sono regolabili dalla dashboard."
    )


def add_operational_notes(doc: Document):
    doc.add_heading("8. Note operative e comportamenti da conoscere", level=1)
    add_bullets(
        doc,
        [
            "Il database usato dall'app è PostgreSQL e viene raggiunto dal servizio web tramite DATABASE_URL.",
            "In Docker il servizio web ascolta sulla porta 5000 e il database è pubblicato sulla porta host 5433.",
            "I container Docker di database e web app sono configurati sul fuso Europe/Rome per allineare log e timestamp operativi all'orario locale italiano.",
            "Per un deploy pubblico HTTPS, impostare APP_ENV=production e SESSION_COOKIE_SECURE=true, mantenendo localhost con SESSION_COOKIE_SECURE=false solo in sviluppo.",
            "Le principali variabili sessione sono SESSION_COOKIE_BROWSER_CLOSE, SESSION_COOKIE_MAX_AGE_SEC e SESSION_INACTIVITY_TIMEOUT_SEC.",
            "L'app supporta anche secret da file con convenzione VAR e VAR_FILE: se esiste un file montato in /run/secrets, il valore del file ha priorità rispetto a quello presente nel .env.",
            "Con Docker Compose locale è possibile mantenere i secret fuori dal repository creando file dedicati sotto docker/secrets e lasciando nel .env solo i riferimenti *_FILE.",
            "Per un deploy pubblico HTTPS, impostare anche SECURITY_ENABLE_HSTS=true dietro reverse proxy TLS stabile.",
            "Il reload automatico del server monitora solo il codice applicativo in app/, così modifiche a manuali o script di supporto non interrompono gli import in corso.",
            "La barra di import può essere numerica o animata, a seconda del tipo di import disponibile.",
            "I log di accesso, import, salvataggio impostazioni ed errori confluiscono nell'audit log.",
            "L'utente standard vede solo i propri percorsi; l'admin vede tutto l'archivio.",
            "Le tabelle routes e auth_connections vengono aggiornate progressivamente in base a login e import.",
            "La dashboard web applica escaping dei contenuti dinamici renderizzati lato browser per ridurre il rischio di XSS sui campi importati dai provider o dai file caricati.",
            "Nota futura: la validazione avanzata della pavimentazione tramite video e AI è prevista come evoluzione più adatta alla futura mobile app, che potrà acquisire video, GPS e contesto del percorso direttamente dal dispositivo.",
        ],
    )

    doc.add_heading("8.1 Condivisione temporanea della web app da PC locale", level=2)
    add_bullets(
        doc,
        [
            "Per mostrare Street Skate a colleghi o tester mantenendo il progetto sul PC locale, è stata validata la pubblicazione temporanea tramite Tailscale Funnel.",
            "La configurazione testata il 14/08/2026 espone l'app locale su Internet senza richiedere Tailscale installato ai client finali.",
            "Prerequisiti: container Docker attivi con web app disponibile su http://localhost:5000 e client Tailscale installato e autenticato sul PC host.",
            "Comando operativo validato: tailscale funnel --bg 5000.",
            "Verifica consigliata: tailscale funnel status deve mostrare il proxy verso http://127.0.0.1:5000.",
            "Il link pubblico generato in dominio ts.net può essere aperto da browser desktop o mobile anche fuori dalla rete locale.",
            "Per interrompere l'esposizione pubblica usare tailscale funnel reset e, se necessario, tailscale serve reset.",
            "Il PC host deve restare acceso e la web app deve rimanere in esecuzione per tutta la durata della demo.",
        ],
    )

    doc.add_heading("8.2 Possibili implementazioni future", level=2)
    add_numbered(
        doc,
        [
            "Usare OSM come baseline statica della pavimentazione tramite i campi surface, smoothness, tracktype e highway.",
            "Introdurre una mappatura rule-based iniziale tra valori OSM e classi funzionali come surface_class, roughness_level e distress_level.",
            "Usare l'AI come livello successivo di arricchimento e validazione, non come sostituto del dato OSM di base.",
            "Separare in modo netto i concetti di materiale della superficie, rugosita e dissestamento, per evitare ambiguità nel ranking dei percorsi.",
            "Estendere l'enrichment con campi strutturati come surface_class, surface_subtype, roughness_level, distress_level, has_cobblestone e relativi livelli di confidence.",
            "Abilitare filtri futuri sui percorsi come pendenza massima assoluta, assenza di sampietrini, rugosita massima e dissestamento massimo, sfruttando i nuovi campi dell'enrichment.",
        ],
    )

    p = doc.add_paragraph(style="ManualCallout")
    p.add_run("Suggerimento operativo: ").bold = True
    p.add_run(
        "se un import sembra fermo, controllare la barra di avanzamento, lo stato testuale in dashboard e, per l'admin, l'audit log o i dati nel database."
    )


def add_quick_reference(doc: Document):
    doc.add_heading("9. Riferimento rapido per ruolo", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_layout(tbl, [1.4, 2.7, 2.4])
    for c, t in zip(tbl.rows[0].cells, ["Ruolo", "Cosa può fare", "Limitazioni principali"]):
        set_cell(c, t, bold=True)
        shade_cell(c, "DCE6F2")
    rows = [
        ("Utente", "Login provider, import provider, filtrare, vedere dettagli, cancellare i propri percorsi", "Non accede alle impostazioni admin e non vede l'archivio completo"),
        ("Admin", "Login locale, vedere tutti i percorsi, import manuale GPX/JSON, filtrare per utente, modificare parametri applicativi, leggere audit log, cancellare qualsiasi percorso", "Non usa il pulsante Importa da provider"),
    ]
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)


def add_score_section(doc: Document):
    doc.add_heading("10. Calcolo del punteggio", level=1)
    doc.add_paragraph(
        "Il punteggio Street Skate viene calcolato automaticamente quando il percorso viene importato o salvato. Una volta scritto sul record, il valore resta fisso e non viene ricalcolato dalla dashboard admin."
    )
    doc.add_heading("10.1 Dati usati nel calcolo", level=2)
    add_bullets(
        doc,
        [
            "Pendenza massima del tratto.",
            "Pendenza massima trattata come valore assoluto, con verso salita/discesa salvato separatamente.",
            "Pendenza media ponderata trattata come valore assoluto.",
            "Tipo di pavimentazione o superficie rilevata.",
            "Pressione atmosferica, se disponibile.",
            "Temperatura, se disponibile.",
            "Velocità massima storica sul tratto, se disponibile.",
        ],
    )
    doc.add_heading("10.2 Logica applicativa", level=2)
    add_bullets(
        doc,
        [
            "L'algoritmo usa pesi base interni all'applicazione.",
            "I pesi non sono modificabili dalla dashboard admin.",
            "I dati mancanti non bloccano il salvataggio del percorso: il sistema usa i valori disponibili al momento dell'import.",
            "Il risultato finale viene normalizzato in uno score su 100 e associato al livello di difficoltà mostrato nelle card.",
        ],
    )
    doc.add_heading("10.3 Effetti pratici in dashboard", level=2)
    add_bullets(
        doc,
        [
            "Lo score appare nelle card percorso e nel riepilogo di difficoltà media.",
            "Il filtro per score massimo usa il valore già salvato sul percorso.",
            "Due utenti che vedono lo stesso percorso salvato vedono lo stesso punteggio registrato per quel record.",
            "Per i percorsi importati da file, l'app può salvare anche una difficoltà dichiarata derivata dal calendario in base alla data nel nome del file, mantenendola separata dal punteggio calcolato.",
        ],
    )


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Street Skate · Manuale utente")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("667085")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)
    doc.add_paragraph()
    add_overview(doc)
    add_login_modes(doc)
    add_screenshots(doc)
    add_user_dashboard(doc)
    add_import_logic(doc)
    add_admin(doc)
    add_admin_settings_detail(doc)
    add_operational_notes(doc)
    add_quick_reference(doc)
    add_score_section(doc)
    for section in doc.sections:
        add_footer(section)
    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        doc.save(FALLBACK_OUT)
        print(FALLBACK_OUT)


if __name__ == "__main__":
    build()
